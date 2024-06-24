from odoo import models, fields, api
import math
from ..util.doc_util import *
from ..util.util import *
from docx import Document
import base64
import os
from ..util.agent_util import *
import statistics
# import sys
# sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
# print(sys.path)
# import translator.Translator_main as ts

# from ..util.translator.Translator_main import *
import io


class task(models.Model):
    _name = 'doc.task'
    type = fields.Selection([('1', '翻译')], string='AI agent', required=True, default='1')
    doc_definition = fields.Many2one('doc.definition', string='PharmaDoc type')
    input_file = fields.Many2one('doc.file', string='Input file')
    output_file = fields.Many2one('doc.file', string='Output file')
    input_lang = fields.Many2one('doc.lang', string='Input language', default=1, readonly=True)
    output_lang = fields.Many2one('doc.lang', string='Output language', default=2, readonly=True)
    status = fields.Selection([('1', '未开始'), ('2', '进行中'), ('3', '已完成')], string='Status', default='1', readonly=True)
    score = fields.Float(string='Task Score', readonly=True)
    queue_task_id = fields.One2many('doc.quene.task', 'task_id', readonly=True)
    msd_usr_id = fields.Char(related='input_file.msd_usr_id', string='MSD ID', readonly=True)
    progress = fields.Float(string='Progress', compute='_compute_progress')

    @api.depends('queue_task_id')
    def _compute_progress(self):
        for rec in self:
            if rec.queue_task_id:
                rec.progress = len(rec.queue_task_id.filtered(lambda x: x.status == '3')) / len(rec.queue_task_id) * 100
            else:
                rec.progress = 0

    def open_input_file_create_form(self):
        # 获取当前记录的 id
        return {
            'type': 'ir.actions.act_window',
            'name': 'Upload source file',
            'res_model': 'doc.file',  # 替换为您的模型名称
            'view_mode': 'form',
            'view_id': self.env.ref('OpenPharmaDoc.doc_file_view_form').id,
            'target': 'new',
            'context': {'default_task_id': self.id, 'default_type': 'input_file'},  # 将当前记录的 id 传递给新建记录的表单窗口
        }

    def open_output_file_create_form(self):
        # 获取当前记录的 id
        return {
            'type': 'ir.actions.act_window',
            'name': 'Upload source file',
            'res_model': 'doc.file',  # 替换为您的模型名称
            'view_mode': 'form',
            'view_id': self.env.ref('OpenPharmaDoc.doc_file_view_form').id,
            'target': 'new',
            'context': {'default_task_id': self.id, 'default_type': 'output_file'},  # 将当前记录的 id 传递给新建记录的表单窗口
        }

    def get_technical_terms(self, term, source_lang, target_lang):
        terms = self.env['doc.translation.technical.terms'].search([('source_term', '=', term), ('source_lang', '=', source_lang.id), ('target_lang', '=', target_lang.id)])
        target_technical_terms = [t.target_term for t in terms]
        return target_technical_terms

    def extract_entity(self, text):
        doc_ontologys = self.doc_definition.doc_ontologys
        source_a_attribute_values = []
        for o in doc_ontologys:
            doc_ontology_attributes = o.doc_ontology_attributes
            input_dict = {
                getLabel(self.source_lang, 'name'): text,
                "a_list": [{"a_id": a.id, "name": a.name} for a in doc_ontology_attributes]
            }
            doc_entity = agent_extract(input_dict)
            for d in doc_entity:
                for a in d:
                    source_a_attribute = a[getLabel(self.source_lang, 'name')]
                    source_a_attribute_values.append(source_a_attribute['value'])
        return source_a_attribute_values

    def action_start(self):
        for rec in self:
            rec['status'] = '2'
            paragraphs = self.env['doc.paragraph'].sudo().search([('doc_file', '=', rec.input_file.id), ('type', 'in', ('1', '2'))])
            # 遍历input文档中的所有段落
            for paragraph in paragraphs:
                if paragraph.text:
                    doc_quene_task = rec.env['doc.quene.task'].sudo().create(
                        {'type': '翻译', 'status': '1', 'source_paragraph': paragraph.id, 'target_lang': rec.output_lang.id, 'source_lang': rec.input_lang.id, 'source_file': rec.input_file.id})
                    self.env.cr.commit()
                    rec.write({'queue_task_id': [(4, doc_quene_task.id, 0)]})


    def get_paragraph_split_num(self, paragraph):
        t_list = []
        p_text = paragraph.text
        idx = 0
        for char in paragraph.text:
            if char == '.' and is_number(paragraph.text[idx - 1]):
                p_text = replace_char_at(p_text, idx, "❤")
            idx += 1
        [t_list.extend(e.replace("❤", ".").split("。")) for e in p_text.split('.')]

        return len([t for t in t_list if t])


    def create_output(self):
        temp_folder_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'temp')
        temp_file_path = os.path.join(temp_folder_path, self.input_file.name)
        binary_data = base64.b64decode(self.input_file.file)
        with open(temp_file_path, 'wb') as f:
            f.write(binary_data)

        # 加载文档
        doc = Document(temp_file_path)
        os.remove(temp_file_path)

        doc_quene_tasks = self.env['doc.quene.task'].sudo().search([("source_file", "=", self.input_file.id),("task_id","=",self.id)])


        # 遍历文档中的所有段落
        pt_idx = 0
        for paragraph in doc.paragraphs:
            if not paragraph.text:
                continue
            target_text = ''
            print(paragraph.text)
            for i in range(0, self.get_paragraph_split_num(paragraph)):
                print(i)
                for t in doc_quene_tasks:
                    if 'p_' + str(pt_idx) == t.source_paragraph_idx:
                        print(t.source_paragraph_idx)
                        target_text += t['target_text']
                        print(target_text)
                        break
                pt_idx += 1
            paragraph.text = target_text

        tt_idx = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if not cell.text:
                        continue
                    print(cell.text)
                    target_text = ''
                    for i in range(0, self.get_paragraph_split_num(cell)):
                        print(i)
                        for t in doc_quene_tasks:
                            if 't_' + str(tt_idx) == t.source_paragraph_idx:
                                print(t.source_paragraph_idx)
                                target_text += t['target_text']
                                print(target_text)
                                break
                        tt_idx += 1
                    cell.text = target_text


        save_path = os.path.join(temp_folder_path, self.input_file.name.replace('.docx', '—Output.docx'))
        # 保存修改后的文档
        doc.save(save_path)

        # 读取docx文件
        with open(save_path, 'rb') as f:
            docx_data = f.read()

        # 将docx文件转换为Base64编码
        docx_base64 = base64.b64encode(docx_data)
        doc_file = self.env['doc.file'].sudo().create({
            'name': self.input_file.name.replace('.docx', '—translate to %s.docx'%(self.output_lang.name)),
            'file': docx_base64,
            'msd_usr_id': self.input_file.msd_usr_id,
            'doc_definition': self.input_file.doc_definition.id,
            'lang': self.output_lang.id
        })

        self.output_file = doc_file.id
        self.status = '3'
        
        scorelist = []
        for t in self.queue_task_id:
            scorelist.append(t['score'])

        self.score = statistics.mean(scorelist)
