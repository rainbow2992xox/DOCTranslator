import math
import json
import time
import docx

from bs4 import BeautifulSoup


def extract_text(html):
    soup = BeautifulSoup(html, 'html.parser')  # 创建BeautifulSoup对象
    text = soup.get_text()  # 提取所有文本内容
    return text


def is_number(s):
    try:
        float(s)
        return True
    except ValueError:
        return False


def replace_char_at(s, index, new_char):
    return s[:index] + new_char + s[index + 1:]


def get_spire_paragraphs(docx_file):
    pass


def get_paragraphs(docx_file):
    doc = docx.Document(docx_file)
    paragraphs = []
    c_title = ''
    title_list = []
    title_index = {
        ' ': []
    }

    for paragraph in doc.paragraphs:
        t_list = []
        if paragraph.style.name.startswith('Heading') and paragraph.text:
            c_title = paragraph.text
            paragraphs.append({'type': 'title', 'title': c_title, 'text': paragraph.text})
            title_list.append(c_title)
            continue

        if paragraph.text:
            # 替换数字中.为 "$&$" 避免被错误分割
            p_text = paragraph.text
            idx = 0
            for char in paragraph.text:
                if char == '.' and is_number(paragraph.text[idx - 1]):
                    p_text = replace_char_at(p_text, idx, "❤")
                idx += 1

            print(p_text)
            [t_list.extend(e.replace("❤", ".").split("。")) for e in p_text.split('.')]
            [paragraphs.append({'title': c_title, 'type': 'text', 'text': t}) for t in t_list if t]

    tb_idx = 0
    c_title = ' '
    for e in doc.element.body.inner_content_elements:
        if "CT_P" in str(e) and e.text in title_list:
            c_title = e.text
            title_index[c_title] = []
        if "CT_Tbl" in str(e):
            title_index[c_title].append(tb_idx)
            tb_idx += 1

    # TODO TABEL加段落
    tb_idx = 0
    for table in doc.tables:
        tb_title = ' '
        for title in title_index:
            if tb_idx in title_index[title]:
                tb_title = title
        row_idx = 1
        for row in table.rows:
            for cell in row.cells:
                t_list = []
                cell_id = 0
                if cell.text:
                    p_text = cell.text
                    print(p_text)
                    idx = 0
                    for char in cell.text:
                        if char == '.' and is_number(cell.text[idx - 1]):
                            p_text = replace_char_at(p_text, idx, "❤")
                        idx += 1

                    [t_list.extend(e.replace("❤", ".").split("。")) for e in p_text.split('.')]
                    title = tb_title + ' 行号：' + str(row_idx) + ' 列号：' + str(cell_id)
                    for t in t_list:
                        if t:
                            paragraphs.append({'title': title, 'type': 'text', 'text': t})
                cell_id += 1
            row_idx += 1

    return paragraphs


def get_tables_index(doc):
    index_list = []
    idx = 0
    for item in doc._element.body.inner_content_elements:
        if "CT_Tbl" in str(item):
            index_list.append(idx)
            idx += 1
    return index_list


def get_tables(docx_file):
    doc = docx.Document(docx_file)
    index_list = get_tables_index(doc)
    tables = []
    idx = 0
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text)
            table_data.append(row_data)
        tables.append((index_list[idx], table_data))
        idx += 1
    return tables


print(get_paragraphs("C:\\Users\\Rainbow\\Desktop\\test.docx"))
